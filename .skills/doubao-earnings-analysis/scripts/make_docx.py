#!/usr/bin/env python3
from __future__ import annotations

import argparse
import re
import zipfile
from pathlib import Path
from xml.sax.saxutils import escape


BASE_DIR = Path(__file__).resolve().parent.parent
DEFAULT_INPUT = BASE_DIR / "report.md"
DEFAULT_OUTPUT = BASE_DIR / "report.docx"

CN_FONT = "宋体-简"
EN_FONT = "Times New Roman"
CITATION_NOTES: dict[str, str] = {}

BLUE = "1F4E79"
TEXT = "222222"
MUTED = "6B7280"
LIGHT_BLUE = "EAF2F8"
LIGHT_GRAY = "F6F8FA"
BORDER = "D9E2EC"


def load_docx_dependency() -> bool:
    """Load python-docx lazily so --help and smoke checks work without it."""
    try:
        from docx import Document as DocxDocument
        from docx.enum.section import WD_ORIENT as DocxWDOrient
        from docx.enum.table import WD_TABLE_ALIGNMENT as DocxWDTableAlignment
        from docx.enum.text import WD_ALIGN_PARAGRAPH as DocxWDAlignParagraph
        from docx.oxml import OxmlElement as DocxOxmlElement
        from docx.oxml.ns import qn as docx_qn
        from docx.opc.constants import RELATIONSHIP_TYPE as DocxRelationshipType
        from docx.shared import Cm as DocxCm
        from docx.shared import Pt as DocxPt
        from docx.shared import RGBColor as DocxRGBColor
    except ModuleNotFoundError as exc:
        if exc.name != "docx":
            raise
        return False

    globals().update(
        {
            "Document": DocxDocument,
            "WD_ORIENT": DocxWDOrient,
            "WD_TABLE_ALIGNMENT": DocxWDTableAlignment,
            "WD_ALIGN_PARAGRAPH": DocxWDAlignParagraph,
            "OxmlElement": DocxOxmlElement,
            "qn": docx_qn,
            "RELATIONSHIP_TYPE": DocxRelationshipType,
            "Cm": DocxCm,
            "Pt": DocxPt,
            "RGBColor": DocxRGBColor,
        }
    )
    return True


def paragraph_xml(text: str, style: str | None = None) -> str:
    style_xml = f'<w:pPr><w:pStyle w:val="{style}"/></w:pPr>' if style else ""
    text_xml = escape(text)
    return f'<w:p>{style_xml}<w:r><w:t xml:space="preserve">{text_xml}</w:t></w:r></w:p>'


def build_basic_docx(input_path: Path, output_path: Path) -> None:
    """Small stdlib fallback for environments without python-docx.

    The full renderer above is used when python-docx is installed. This fallback
    keeps delivery smoke tests useful by producing a valid, plain DOCX instead
    of failing at import time.
    """
    paragraphs: list[str] = []
    for raw in input_path.read_text(encoding="utf-8").splitlines():
        line = raw.strip()
        if not line:
            paragraphs.append(paragraph_xml(""))
        elif line.startswith("# "):
            paragraphs.append(paragraph_xml(line[2:].strip(), "Title"))
        elif line.startswith("## "):
            paragraphs.append(paragraph_xml(line[3:].strip(), "Heading1"))
        elif line.startswith("### "):
            paragraphs.append(paragraph_xml(line[4:].strip(), "Heading2"))
        elif re.match(r"^[-*]\s+", line):
            paragraphs.append(paragraph_xml("• " + re.sub(r"^[-*]\s+", "", line)))
        else:
            paragraphs.append(paragraph_xml(line))

    document_xml = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
  <w:body>
    {body}
    <w:sectPr><w:pgSz w:w="11906" w:h="16838"/><w:pgMar w:top="1440" w:right="1440" w:bottom="1440" w:left="1440"/></w:sectPr>
  </w:body>
</w:document>
""".format(body="\n    ".join(paragraphs))
    content_types = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">
  <Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>
  <Default Extension="xml" ContentType="application/xml"/>
  <Override PartName="/word/document.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>
</Types>
"""
    rels = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
  <Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="word/document.xml"/>
</Relationships>
"""
    output_path.parent.mkdir(parents=True, exist_ok=True)
    with zipfile.ZipFile(output_path, "w", compression=zipfile.ZIP_DEFLATED) as docx:
        docx.writestr("[Content_Types].xml", content_types)
        docx.writestr("_rels/.rels", rels)
        docx.writestr("word/document.xml", document_xml)
    print("[警告] 未安装 python-docx，已使用简易 DOCX 兜底渲染；如需完整版式请安装 python-docx。")
    print(f"DOCX saved: {output_path}")


def set_rfonts(element, cn_font: str = CN_FONT, en_font: str = EN_FONT) -> None:
    r_pr = element.rPr
    if r_pr is None:
        r_pr = element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:ascii"), en_font)
    r_fonts.set(qn("w:hAnsi"), en_font)
    r_fonts.set(qn("w:cs"), en_font)
    r_fonts.set(qn("w:eastAsia"), cn_font)


def set_style_font(style, size: float, bold: bool = False, color: str = TEXT) -> None:
    font = style.font
    font.name = EN_FONT
    font.size = Pt(size)
    font.bold = bold
    font.color.rgb = RGBColor.from_string(color)
    set_rfonts(style.element)


def set_run_font(run, size: float | None = None, bold: bool | None = None, color: str | None = None) -> None:
    run.font.name = EN_FONT
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    set_rfonts(run._element)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    for child in tc_pr.findall(qn("w:shd")):
        tc_pr.remove(child)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_vertical_center(cell) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    v_align = tc_pr.find(qn("w:vAlign"))
    if v_align is None:
        v_align = OxmlElement("w:vAlign")
        tc_pr.append(v_align)
    v_align.set(qn("w:val"), "center")


def set_cell_margins(cell, top: int = 80, start: int = 100, bottom: int = 80, end: int = 100) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color: str = BORDER) -> None:
    tbl_pr = table._tbl.tblPr
    existing = tbl_pr.find(qn("w:tblBorders"))
    if existing is not None:
        tbl_pr.remove(existing)
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = OxmlElement(f"w:{edge}")
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), "4")
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), color)
        borders.append(tag)
    tbl_pr.append(borders)


def set_table_cell_paragraph_format(paragraph) -> None:
    paragraph.paragraph_format.space_before = Pt(2)
    paragraph.paragraph_format.space_after = Pt(2)
    paragraph.paragraph_format.line_spacing = 1.35


def add_bottom_border(paragraph, color: str = BLUE, size: str = "8") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def add_paragraph_shading(paragraph, fill: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("第 ")
    set_run_font(run, size=9, color=MUTED)

    def append_field_node(node) -> None:
        field_run = OxmlElement("w:r")
        field_run.append(node)
        paragraph._p.append(field_run)

    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    append_field_node(fld_begin)

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    append_field_node(instr)

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    append_field_node(fld_end)

    run = paragraph.add_run(" 页")
    set_run_font(run, size=9, color=MUTED)


def add_hyperlink(paragraph, text: str, url: str, bold: bool = False) -> None:
    part = paragraph.part
    r_id = part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    run.append(r_pr)

    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), EN_FONT)
    r_fonts.set(qn("w:hAnsi"), EN_FONT)
    r_fonts.set(qn("w:cs"), EN_FONT)
    r_fonts.set(qn("w:eastAsia"), CN_FONT)
    r_pr.append(r_fonts)

    color = OxmlElement("w:color")
    color.set(qn("w:val"), BLUE)
    r_pr.append(color)

    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)

    if bold:
        r_pr.append(OxmlElement("w:b"))

    node = OxmlElement("w:t")
    node.text = text
    run.append(node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_citation_reference(paragraph, number: str) -> None:
    note = CITATION_NOTES.get(number)
    if not note:
        run = paragraph.add_run(f"[{number}]")
        set_run_font(run, size=8.5, bold=True, color=BLUE)
        run.font.superscript = True
        return

    def append_field_char(char_type: str) -> None:
        field_run = OxmlElement("w:r")
        field_char = OxmlElement("w:fldChar")
        field_char.set(qn("w:fldCharType"), char_type)
        field_run.append(field_char)
        paragraph._p.append(field_run)

    append_field_char("begin")

    tooltip = f"来源 {number}: {note[:230]}".replace('"', "'")
    instr_run = OxmlElement("w:r")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f' HYPERLINK \\l "source_{number}" \\o "{tooltip}" '
    instr_run.append(instr)
    paragraph._p.append(instr_run)

    append_field_char("separate")

    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    run.append(r_pr)

    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), EN_FONT)
    r_fonts.set(qn("w:hAnsi"), EN_FONT)
    r_fonts.set(qn("w:cs"), EN_FONT)
    r_fonts.set(qn("w:eastAsia"), CN_FONT)
    r_pr.append(r_fonts)

    color = OxmlElement("w:color")
    color.set(qn("w:val"), BLUE)
    r_pr.append(color)

    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "none")
    r_pr.append(underline)

    bold = OxmlElement("w:b")
    r_pr.append(bold)

    size = OxmlElement("w:sz")
    size.set(qn("w:val"), "17")
    r_pr.append(size)

    vert = OxmlElement("w:vertAlign")
    vert.set(qn("w:val"), "superscript")
    r_pr.append(vert)

    node = OxmlElement("w:t")
    node.text = f"[{number}]"
    run.append(node)
    paragraph._p.append(run)

    append_field_char("end")


def add_bookmark_start(paragraph, name: str, bookmark_id: int) -> None:
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bookmark_id))
    start.set(qn("w:name"), name)
    paragraph._p.append(start)


def add_bookmark_end(paragraph, bookmark_id: int) -> None:
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bookmark_id))
    paragraph._p.append(end)


def clean_text(text: str) -> str:
    text = text.strip()
    if text.startswith("**") and text.count("**") == 1:
        text = text[2:].strip()
    return normalize_corner_quotes(text)


def normalize_corner_quotes(text: str) -> str:
    """Use corner quotes as draft markup, then render standard Chinese quotes."""
    text = re.sub(r"「([^」]*)」", r"“\1”", text)
    text = re.sub(r"『([^』]*)』", r"‘\1’", text)
    return text


def strip_markdown_for_tooltip(text: str) -> str:
    text = re.sub(r"^\[\^\d+\]:?\s*", "", text)
    text = re.sub(r"\[([^\]]+)\]\(([^)]+)\)", r"\1", text)
    text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
    return normalize_corner_quotes(text.strip())


def strip_leading_footnote_marker(text: str) -> str:
    return re.sub(r"^\[\^\d+\]\s*", "", text.strip())


def extract_citation_notes(lines: list[str]) -> dict[str, str]:
    notes: dict[str, str] = {}
    in_notes = False
    for raw in lines:
        line = raw.strip()
        if line == "文中引用对应以下来源：":
            in_notes = True
            continue
        if not in_notes:
            # Also match [^N]: ... footnote definitions anywhere in the file
            fn_match = re.match(r"^\[\^(\d+)\]:\s*(.+)$", line)
            if fn_match:
                notes[fn_match.group(1)] = strip_markdown_for_tooltip(fn_match.group(2))
            continue
        if line.startswith("- **未获取清单**") or line.startswith("## "):
            break
        match = re.match(r"^(\d+)\.\s+(.+)$", line)
        if match:
            notes[match.group(1)] = strip_markdown_for_tooltip(match.group(2))
    return notes


INLINE_RE = re.compile(r"(\[\^([0-9]+)\]|\[([^\]]+)\]\(([^)]+)\)|\*\*([^*]+)\*\*)")
FACT_REF_RE = re.compile(r"\{facts?:[^{}\n]*\}")


def add_inline(paragraph, text: str, size: float | None = None, default_bold: bool = False) -> None:
    text = clean_text(text)
    text = FACT_REF_RE.sub("", text)
    pos = 0
    for match in INLINE_RE.finditer(text):
        if match.start() > pos:
            run = paragraph.add_run(text[pos : match.start()])
            set_run_font(run, size=size, bold=default_bold)
        if match.group(2) is not None:
            add_citation_reference(paragraph, match.group(2))
        elif match.group(3) is not None:
            add_hyperlink(paragraph, match.group(3), match.group(4), bold=default_bold)
        else:
            bold_content = match.group(5)
            sub_pos = 0
            for sub_match in re.finditer(r"\[\^([0-9]+)\]", bold_content):
                if sub_match.start() > sub_pos:
                    run = paragraph.add_run(bold_content[sub_pos:sub_match.start()])
                    set_run_font(run, size=size, bold=True, color=TEXT)
                add_citation_reference(paragraph, sub_match.group(1))
                sub_pos = sub_match.end()
            if sub_pos < len(bold_content):
                run = paragraph.add_run(bold_content[sub_pos:])
                set_run_font(run, size=size, bold=True, color=TEXT)
        pos = match.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        set_run_font(run, size=size, bold=default_bold)


def parse_table_row(line: str) -> list[str]:
    line = FACT_REF_RE.sub("", line)
    return [cell.strip() for cell in line.strip().strip("|").split("|")]


def parse_alignment(line: str) -> list[str]:
    alignments = []
    for cell in parse_table_row(line):
        left = cell.startswith(":")
        right = cell.endswith(":")
        if left and right:
            alignments.append("center")
        elif right:
            alignments.append("right")
        else:
            alignments.append("left")
    return alignments


MISSING_TABLE_VALUE_RE = re.compile(r"^(?:[-—–]|N/?A|NA|n\.a\.|未获取|不适用|无)?$", re.IGNORECASE)
NUMERIC_TABLE_VALUE_RE = re.compile(
    r"^[+−-]?\s*(?:\d{1,3}(?:,\d{3})+|\d+)(?:\.\d+)?\s*"
    r"(?:%|pct|pcts|个百分点|倍|x|X|亿|万|元|亿元|亿美元|美元|人民币|港元)?$",
    re.IGNORECASE,
)


def strip_table_cell_markup(value: str) -> str:
    value = normalize_corner_quotes(value.strip())
    value = FACT_REF_RE.sub("", value)
    value = re.sub(r"\[\^\d+\]", "", value)
    value = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", value)
    value = re.sub(r"\*\*([^*]+)\*\*", r"\1", value)
    return value.strip()


def is_missing_table_value(value: str) -> bool:
    return bool(MISSING_TABLE_VALUE_RE.fullmatch(strip_table_cell_markup(value)))


def is_numeric_table_value(value: str) -> bool:
    value = strip_table_cell_markup(value)
    value = value.replace("−", "-").replace("＋", "+")
    value = re.sub(r"\s+", " ", value)
    return bool(NUMERIC_TABLE_VALUE_RE.fullmatch(value))


def infer_numeric_column_alignments(rows: list[list[str]], col_count: int, alignments: list[str]) -> list[str]:
    inferred = [alignments[i] if i < len(alignments) else "left" for i in range(col_count)]
    for col_index in range(col_count):
        values = [row[col_index] if col_index < len(row) else "" for row in rows]
        non_missing = [value for value in values if not is_missing_table_value(value)]
        if non_missing and all(is_numeric_table_value(value) for value in non_missing):
            inferred[col_index] = "center"
    return inferred


def infer_table_alignments(
    rows: list[list[str]], col_count: int, markdown_alignments: list[str], table_index: int
) -> list[str]:
    if table_index == 0:
        return ["left"] + ["center"] * max(col_count - 1, 0)
    return infer_numeric_column_alignments(rows, col_count, markdown_alignments)


def add_table(doc: Document, table_lines: list[str], table_index: int = 0) -> None:
    headers = parse_table_row(table_lines[0])
    markdown_alignments = parse_alignment(table_lines[1])
    rows = [parse_table_row(line) for line in table_lines[2:]]
    col_count = len(headers)
    alignments = infer_table_alignments(rows, col_count, markdown_alignments, table_index)

    table = doc.add_table(rows=1, cols=col_count)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    set_table_borders(table)

    header_cells = table.rows[0].cells
    for i, text in enumerate(headers):
        cell = header_cells[i]
        set_cell_vertical_center(cell)
        set_cell_shading(cell, BLUE)
        set_cell_margins(cell)
        paragraph = cell.paragraphs[0]
        set_table_cell_paragraph_format(paragraph)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_inline(paragraph, text, size=9.5, default_bold=True)
        for run in paragraph.runs:
            set_run_font(run, size=9.5, bold=True, color="FFFFFF")

    tr_pr = table.rows[0]._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)

    table_font_size = 8.5 if col_count <= 3 else 8.8
    for row_index, row in enumerate(rows, start=1):
        cells = table.add_row().cells
        if row_index % 2 == 0:
            for cell in cells:
                set_cell_shading(cell, LIGHT_GRAY)
        for col_index in range(col_count):
            cell = cells[col_index]
            set_cell_vertical_center(cell)
            set_cell_margins(cell)
            paragraph = cell.paragraphs[0]
            set_table_cell_paragraph_format(paragraph)
            alignment = alignments[col_index] if col_index < len(alignments) else "left"
            if alignment == "right":
                paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            elif alignment == "center":
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            value = row[col_index] if col_index < len(row) else ""
            add_inline(paragraph, value, size=table_font_size)

    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(3)


def is_table_separator(line: str) -> bool:
    return bool(re.match(r"^\s*\|?\s*:?-{3,}:?\s*(\|\s*:?-{3,}:?\s*)+\|?\s*$", line))


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.1)
    section.right_margin = Cm(2.1)
    section.footer_distance = Cm(0.9)

    styles = doc.styles
    set_style_font(styles["Normal"], 10.5, color=TEXT)
    styles["Normal"].paragraph_format.line_spacing = 1.35
    styles["Normal"].paragraph_format.space_after = Pt(5)

    for style_name in ("List Bullet", "List Bullet 2", "List Bullet 3", "List Number"):
        try:
            style = styles[style_name]
        except KeyError:
            continue
        set_style_font(style, 10.5, color=TEXT)
        style.paragraph_format.space_after = Pt(5)
        style.paragraph_format.line_spacing = 1.32
    styles["List Number"].paragraph_format.left_indent = Cm(0.75)
    styles["List Number"].paragraph_format.first_line_indent = Cm(-0.45)

    set_style_font(styles["Title"], 21, bold=True, color=BLUE)
    styles["Title"].paragraph_format.space_after = Pt(7)
    styles["Title"].paragraph_format.line_spacing = 1.12

    set_style_font(styles["Subtitle"], 10, color=MUTED)
    styles["Subtitle"].paragraph_format.space_after = Pt(9)

    set_style_font(styles["Heading 1"], 14, bold=True, color=BLUE)
    styles["Heading 1"].paragraph_format.space_before = Pt(14)
    styles["Heading 1"].paragraph_format.space_after = Pt(7)
    styles["Heading 1"].paragraph_format.keep_with_next = True

    set_style_font(styles["Heading 2"], 12, bold=True, color=TEXT)
    styles["Heading 2"].paragraph_format.space_before = Pt(10)
    styles["Heading 2"].paragraph_format.space_after = Pt(5)
    styles["Heading 2"].paragraph_format.keep_with_next = True

    set_style_font(styles["Heading 3"], 10.5, bold=True, color=TEXT)
    styles["Heading 3"].paragraph_format.space_before = Pt(8)
    styles["Heading 3"].paragraph_format.space_after = Pt(4)
    styles["Heading 3"].paragraph_format.keep_with_next = True

    footer = section.footer.paragraphs[0]
    add_page_number(footer)


def add_rule(doc: Document) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(1)
    paragraph.paragraph_format.space_after = Pt(9)
    add_bottom_border(paragraph)


TABLE_NOTE_RE = re.compile(
    r"^(数据来源|口径提示|注[：:]|备注[：:])"
)


def is_table_note(text: str) -> bool:
    return bool(TABLE_NOTE_RE.match(text.lstrip("*").strip()))


def add_table_note_paragraph(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(1)
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.paragraph_format.line_spacing = 1.12
    add_inline(paragraph, text, size=8.5)
    for run in paragraph.runs:
        set_run_font(run, size=8.5, color=MUTED)


def add_table_note_bullet(doc: Document, text: str, level: int = 0) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Cm(0.42 + 0.32 * level)
    paragraph.paragraph_format.first_line_indent = Cm(-0.18)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(2)
    paragraph.paragraph_format.line_spacing = 1.12
    run = paragraph.add_run("• ")
    set_run_font(run, size=8.5, color=MUTED)
    add_inline(paragraph, text, size=8.5)
    for run in paragraph.runs:
        set_run_font(run, size=8.5, color=MUTED)


def add_table_note_numbered(doc: Document, number: str, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Cm(0.48)
    paragraph.paragraph_format.first_line_indent = Cm(-0.24)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(2)
    paragraph.paragraph_format.line_spacing = 1.12
    run = paragraph.add_run(f"{number}. ")
    set_run_font(run, size=8.5, color=MUTED)
    add_inline(paragraph, text, size=8.5)
    for run in paragraph.runs:
        set_run_font(run, size=8.5, color=MUTED)


def add_body_paragraph(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    add_inline(paragraph, text)

    if text.startswith("**判定：") or text.startswith("判定："):
        add_paragraph_shading(paragraph, LIGHT_BLUE)
        paragraph.paragraph_format.left_indent = Cm(0.15)
        paragraph.paragraph_format.right_indent = Cm(0.15)
        paragraph.paragraph_format.space_before = Pt(2)
        paragraph.paragraph_format.space_after = Pt(7)
        for run in paragraph.runs:
            set_run_font(run, bold=True, color=BLUE)


def add_bullet_paragraph(doc: Document, text: str, level: int = 0) -> None:
    style_name = "List Bullet" if level <= 0 else f"List Bullet {min(level + 1, 3)}"
    try:
        paragraph = doc.add_paragraph(style=style_name)
    except KeyError:
        paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    add_inline(paragraph, text)


def add_numbered_paragraph(doc: Document, number: str, text: str) -> None:
    """Render a numbered list item manually with hanging indent (avoids Word numbering bugs)."""
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Cm(0.75)
    paragraph.paragraph_format.first_line_indent = Cm(-0.45)
    paragraph.paragraph_format.space_after = Pt(5)
    paragraph.paragraph_format.line_spacing = 1.32
    run = paragraph.add_run(f"{number}. ")
    set_run_font(run, size=10.5)
    add_inline(paragraph, text, size=10.5)


def build_docx(input_path: Path = DEFAULT_INPUT, output_path: Path = DEFAULT_OUTPUT) -> None:
    if not load_docx_dependency():
        build_basic_docx(input_path, output_path)
        return

    global CITATION_NOTES
    doc = Document()
    configure_document(doc)

    lines = input_path.read_text(encoding="utf-8").splitlines()
    CITATION_NOTES = extract_citation_notes(lines)
    i = 0
    seen_title = False
    in_source_notes = False
    prev_was_numbered = False
    after_table = False
    in_table_note = False
    table_count = 0

    while i < len(lines):
        raw = lines[i]
        line = raw.strip()
        if not line:
            # A blank line ends a table-note block. Keep after_table so a note
            # can still follow a table after one blank line.
            in_table_note = False
            i += 1
            continue
        if line == "文中引用对应以下来源：":
            in_source_notes = True

        if line == "---":
            add_rule(doc)
            prev_was_numbered = False
            after_table = False
            in_table_note = False
            i += 1
            continue

        if line.startswith("# "):
            after_table = False
            in_table_note = False
            paragraph = doc.add_paragraph(style="Title" if not seen_title else "Heading 1")
            add_inline(paragraph, line[2:].strip(), size=21 if not seen_title else 14, default_bold=True)
            seen_title = True
            prev_was_numbered = False
            i += 1
            continue

        if seen_title and line.startswith("**") and "｜" in line and "业绩点评" in line:
            after_table = False
            in_table_note = False
            paragraph = doc.add_paragraph(style="Subtitle")
            add_inline(paragraph, line, size=10)
            i += 1
            continue

        if line.startswith("## "):
            after_table = False
            in_table_note = False
            heading_text = line[3:].strip()
            numbered_heading = re.match(r"^(\d+)\.\s+(.+)$", heading_text)
            paragraph = doc.add_paragraph(style="Heading 1")
            if numbered_heading:
                run = paragraph.add_run(f"{numbered_heading.group(1)}. ")
                set_run_font(run, size=14, bold=True, color=BLUE)
                add_inline(paragraph, numbered_heading.group(2), size=14, default_bold=True)
            else:
                add_inline(paragraph, heading_text, size=14, default_bold=True)
            prev_was_numbered = False
            i += 1
            continue

        if line.startswith("#### "):
            after_table = False
            in_table_note = False
            paragraph = doc.add_paragraph(style="Heading 3")
            add_inline(paragraph, line[5:].strip(), size=10.5, default_bold=True)
            prev_was_numbered = False
            i += 1
            continue

        if line.startswith("### "):
            after_table = False
            in_table_note = False
            heading_text = line[4:].strip()
            numbered_heading = re.match(r"^(\d+)\.\s+(.+)$", heading_text)
            paragraph = doc.add_paragraph(style="Heading 2")
            if numbered_heading:
                run = paragraph.add_run(f"{numbered_heading.group(1)}. ")
                set_run_font(run, size=12, bold=True, color=TEXT)
                add_inline(paragraph, numbered_heading.group(2), size=12, default_bold=True)
            else:
                add_inline(paragraph, heading_text, size=12, default_bold=True)
            prev_was_numbered = False
            i += 1
            continue

        if (
            line.startswith("|")
            and i + 1 < len(lines)
            and is_table_separator(lines[i + 1].strip())
        ):
            table_lines = [line, lines[i + 1].strip()]
            i += 2
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i].strip())
                i += 1
            add_table(doc, table_lines, table_index=table_count)
            table_count += 1
            after_table = True
            in_table_note = False
            continue

        if in_source_notes and line.startswith("- **未获取清单**"):
            in_source_notes = False

        # Markdown footnote definitions: [^N]: source text
        fn_def = re.match(r"^\[\^(\d+)\]:\s*(.+)$", line)
        if fn_def:
            after_table = False
            in_table_note = False
            fn_number = fn_def.group(1)
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.left_indent = Cm(0.35)
            paragraph.paragraph_format.first_line_indent = Cm(-0.35)
            add_bookmark_start(paragraph, f"source_{fn_number}", 1000 + int(fn_number))
            run = paragraph.add_run(f"[{fn_number}] ")
            set_run_font(run, size=9, color=MUTED)
            add_inline(paragraph, fn_def.group(2), size=9)
            add_bookmark_end(paragraph, 1000 + int(fn_number))
            prev_was_numbered = False
            i += 1
            continue

        raw_bullet = re.match(r"^(\s*)[-*]\s+(.+)$", raw)
        bullet = re.match(r"^[-*]\s+(.+)$", line)
        if raw_bullet:
            level = len(raw_bullet.group(1).replace("\t", "    ")) // 2
            if in_table_note:
                add_table_note_bullet(doc, raw_bullet.group(2), level=level)
            else:
                add_bullet_paragraph(doc, raw_bullet.group(2), level=level)
            after_table = False
            prev_was_numbered = False
            i += 1
            continue

        if bullet:
            if in_table_note:
                add_table_note_bullet(doc, bullet.group(1))
            else:
                add_bullet_paragraph(doc, bullet.group(1))
            after_table = False
            prev_was_numbered = False
            i += 1
            continue

        numbered = re.match(r"^(\d+)\.\s+(.+)$", line)
        if numbered:
            number = numbered.group(1)
            if in_table_note:
                add_table_note_numbered(doc, number, numbered.group(2))
                prev_was_numbered = False
            elif in_source_notes and number in CITATION_NOTES:
                paragraph = doc.add_paragraph()
                paragraph.paragraph_format.left_indent = Cm(0.35)
                paragraph.paragraph_format.first_line_indent = Cm(-0.35)
                add_bookmark_start(paragraph, f"source_{number}", 1000 + int(number))
                run = paragraph.add_run(f"{number}. ")
                set_run_font(run)
                add_inline(paragraph, strip_leading_footnote_marker(numbered.group(2)))
                add_bookmark_end(paragraph, 1000 + int(number))
                prev_was_numbered = False
            else:
                add_numbered_paragraph(doc, number, numbered.group(2))
                prev_was_numbered = True
            after_table = False
            i += 1
            continue

        prev_was_numbered = False
        if after_table and is_table_note(line):
            add_table_note_paragraph(doc, line)
            after_table = False
            in_table_note = True
        else:
            in_table_note = False
            add_body_paragraph(doc, line)
            after_table = False
        i += 1

    doc.core_properties.title = ""
    doc.core_properties.subject = "季度业绩点评"
    doc.core_properties.keywords = ""
    doc.save(output_path)
    print(f"DOCX saved: {output_path}")


if __name__ == "__main__":
    parser = argparse.ArgumentParser(
        description="Internal helper. Prefer finalize_report.py for delivery; this only converts source markdown to DOCX."
    )
    parser.add_argument("input", nargs="?", default=str(DEFAULT_INPUT), help="Input markdown file")
    parser.add_argument("-o", "--output", default=None, help="Output .docx path (default: same dir as input, .docx extension)")
    args = parser.parse_args()
    input_path = Path(args.input).resolve()
    if input_path.name.endswith("-display.md"):
        raise SystemExit(
            "[错误] 输入文件看起来是 display markdown。"
            "请运行 finalize_report.py，并传入保留 {fact:...} 绑定的源 markdown。"
        )
    if args.output:
        output = Path(args.output).resolve()
    else:
        output = input_path.with_suffix(".docx")
    build_docx(input_path, output)
